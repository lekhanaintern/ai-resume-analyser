import PyPDF2
import docx
import pdfplumber
import re


# ─────────────────────────────────────────────────────────────────────────────
#  PARAGRAPH → BULLET CONVERTER
#  Runs on all extracted text (PDF + DOCX) before it reaches the rest of app.
#  Splits paragraph blobs into one bullet per sentence.
# ─────────────────────────────────────────────────────────────────────────────

_HEADER_RE  = re.compile(r'^[A-Z][A-Z\s/&\-]{2,}$')          # ALL-CAPS header
_SEP_RE     = re.compile(r'^[\-─═]{3,}$')                      # separator lines
_BULLET_RE  = re.compile(r'^[\-\•\*\►\▸\–\—\>]')              # already a bullet
_CONTACT_RE = re.compile(                                       # name/contact rows
    r'@|\+?\d[\d\s\-(). ]{5,}\d|linkedin|github|http|www|\|', re.I)
_DATE_RE    = re.compile(                                       # job title / date lines
    r'\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|'
    r'january|february|march|april|june|july|august|september|october|november|december|'
    r'20\d\d|19\d\d|present|current|till|to)\b', re.I)

def _sentences(text: str):
    """Split a blob into individual sentences."""
    # Split at . ! ? followed by whitespace + capital letter
    parts = re.split(r'(?<=[.!?])\s+(?=[A-Z])', text.strip())
    out = []
    for p in parts:
        p = p.strip()
        if not p:
            continue
        # Further split at semicolons if each part is substantial
        sub = re.split(r';\s+', p)
        if len(sub) > 1 and all(len(s.split()) >= 4 for s in sub):
            out.extend(sub)
        else:
            out.append(p)
    return out


def paragraphs_to_bullets(raw_text: str) -> str:
    """
    Convert paragraph-style resume text to bullet-point format.

    Rules:
    - Lines already starting with a bullet: kept as-is
    - ALL-CAPS section headers: kept as-is
    - Separator lines (--- ═══): kept as-is
    - Contact / URL lines: kept as-is
    - Short lines (≤ 6 words): kept as-is (job titles, dates, names)
    - Everything else (7+ words): split into sentences → one bullet each
    """
    lines    = raw_text.split('\n')
    out      = []
    line_num = 0

    for line in lines:
        s     = line.strip()
        words = s.split()

        # Blank line
        if not s:
            out.append('')
            line_num += 1
            continue

        # Already a bullet
        if _BULLET_RE.match(s):
            out.append(line)
            line_num += 1
            continue

        # Section header (ALL CAPS, short)
        if _HEADER_RE.match(s) and len(words) <= 7:
            out.append(line)
            line_num += 1
            continue

        # Separator
        if _SEP_RE.match(s):
            out.append(line)
            line_num += 1
            continue

        # Contact / URL lines — always skip
        if _CONTACT_RE.search(s):
            out.append(line)
            line_num += 1
            continue

        # Short lines (job titles, dates, names) — skip if ≤ 12 words
        if len(words) <= 12 and _DATE_RE.search(s):
            out.append(line)
            line_num += 1
            continue

        # Very short lines (≤ 6 words) — keep as-is regardless
        if len(words) <= 6:
            out.append(line)
            line_num += 1
            continue

        # ── Paragraph blob → bullets ──────────────────────────────────
        sents = _sentences(s)

        if len(sents) <= 1:
            # Single sentence or no split found — still make it a bullet
            out.append('- ' + s)
        else:
            for sent in sents:
                sent = sent.strip().rstrip('.')
                if len(sent.split()) >= 4:
                    out.append('- ' + sent + '.')
                else:
                    out.append(sent)

        line_num += 1

    return '\n'.join(out)


def extract_text_from_pdf(file):
    text = ""
    try:
        file.seek(0)
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        if text.strip():
            return paragraphs_to_bullets(text[:50000])
    except Exception as e:
        print(f"[PDF] pdfplumber failed: {e}")

    try:
        file.seek(0)
        pdf_reader = PyPDF2.PdfReader(file)
        for page_num in range(min(len(pdf_reader.pages), 50)):
            try:
                page_text = pdf_reader.pages[page_num].extract_text()
                if page_text:
                    text += page_text + "\n"
            except Exception as e:
                print(f"[PDF] PyPDF2 page {page_num+1} warning: {e}")
        if text.strip():
            return paragraphs_to_bullets(text[:50000])
    except Exception as e:
        print(f"[PDF] PyPDF2 failed: {e}")

    try:
        import pytesseract
        from pdf2image import convert_from_bytes
        pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        poppler_path = r'C:\poppler\poppler-25.12.0\Library\bin'
        file.seek(0)
        images = convert_from_bytes(file.read(), dpi=200, poppler_path=poppler_path)
        for image in images:
            text += pytesseract.image_to_string(image) + "\n"
        if text.strip():
            return paragraphs_to_bullets(text[:50000])
    except Exception as e:
        print(f"[PDF] Tesseract failed: {e}")

    return text


def extract_text_from_docx(file):
    try:
        doc = docx.Document(file)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text.strip()]
                if row_text:
                    tables_text.append(' '.join(row_text))
        text = '\n'.join(paragraphs + tables_text)
        if not text.strip():
            raise Exception("No text could be extracted from DOCX file.")
        return paragraphs_to_bullets(text[:50000])
    except Exception as e:
        raise Exception(f"Error reading DOCX: {str(e)}")